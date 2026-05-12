import streamlit as st
import google.generativeai as genai
import json
import os
from datetime import datetime
from io import BytesIO
from docx import Document

# --- Налаштування Gemini ---
# Перевіряємо різні варіанти назви ключа. 
# У Streamlit Secrets зазвичай додають GEMINI_API_KEY.
api_key = st.secrets.get("GEMINI_API_KEY") or os.environ.get("GEMINI_API_KEY")

if api_key:
    # Важливо: конфігуруємо API перед використанням
    genai.configure(api_key=api_key)
    # Використовуємо modern flash модель
    model = genai.GenerativeModel('gemini-1.5-flash')
else:
    api_key = None

# --- Функції Експорту ---
def export_to_docx(decision):
    doc = Document()
    doc.add_heading('Аналітичний звіт про рішення', 0)
    
    doc.add_heading('Запит', level=1)
    doc.add_paragraph(decision['query'])
    doc.add_paragraph(f"Дата проведення аналізу: {decision['timestamp']}")
    
    doc.add_heading('Переваги та Недоліки', level=1)
    doc.add_heading('Переваги:', level=2)
    for p in decision['analysis']['prosCons']['pros']:
        doc.add_paragraph(f"• {p}")
    
    doc.add_heading('Недоліки:', level=2)
    for c in decision['analysis']['prosCons']['cons']:
        doc.add_paragraph(f"• {c}")
    
    doc.add_heading('SWOT Аналіз', level=1)
    swot = decision['analysis']['swot']
    for label, items in [("Сильні сторони (Strengths)", swot['strengths']), 
                         ("Слабкі сторони (Weaknesses)", swot['weaknesses']), 
                         ("Можливості (Opportunities)", swot['opportunities']), 
                         ("Загрози (Threats)", swot['threats'])]:
        doc.add_heading(label, level=2)
        for item in items:
            doc.add_paragraph(f"• {item}")
            
    doc.add_heading('Підсумок та рекомендація', level=1)
    doc.add_paragraph(decision['analysis']['summary'])
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Логіка AI ---
def get_ai_analysis(query):
    # Промпт з чіткими інструкціями щодо JSON
    prompt = f"""
    Проаналізуй наступне рішення або дилему: "{query}"
    Надай структуровану відповідь українською мовою.
    
    Формат відповіді має бути СУВОРИМ JSON з наступною структурою:
    {{
      "prosCons": {{
        "pros": ["аргумент за 1", "аргумент за 2", "аргумент за 3"],
        "cons": ["аргумент проти 1", "аргумент проти 2", "аргумент проти 3"]
      }},
      "swot": {{
        "strengths": ["сила 1", "сила 2"],
        "weaknesses": ["слабкість 1", "слабкість 2"],
        "opportunities": ["можливість 1"],
        "threats": ["ризик 1"]
      }},
      "summary": "Короткий фінальний підсумок та рекомендація."
    }}
    
    Не додавай жодних пояснень поза JSON.
    """
    
    try:
        # Використовуємо режим генерації JSON для надійності
        response = model.generate_content(
            prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        
        if not response.text:
            return None
            
        return json.loads(response.text)
    except Exception as e:
        # Для налагодження у консолі Streamlit
        st.error(f"Помилка аналізу: {str(e)}")
        return None

# --- UI Налаштування ---
st.set_page_config(page_title="Аналізатор рішень", page_icon="🧠", layout="centered")

# Стилізація
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button { border-radius: 10px; height: 3em; font-weight: bold; }
    .stTextArea>div>div>textarea { border-radius: 15px; }
    </style>
    """, unsafe_allow_html=True)

if "history" not in st.session_state:
    st.session_state.history = []

st.title("🧠 Аналізатор рішень")
st.write("Ваш персональний ШІ-помічник для виваженого вибору.")

if not api_key:
    st.info("ℹ️ Потрібно налаштувати API ключ. Додайте `GEMINI_API_KEY` у Secrets вашого Streamlit додатку.")
else:
    # Секція введення
    with st.container():
        query = st.text_area("Опишіть ваше рішення чи ситуацію:", height=100, placeholder="Наприклад: Чи варто мені переходити на фріланс зараз?")
        
        if st.button("Аналізувати ситуацію 🔍", use_container_width=True):
            if query:
                with st.spinner("ШІ аналізує варіанти..."):
                    analysis_result = get_ai_analysis(query)
                    if analysis_result:
                        new_decision = {
                            "query": query,
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "analysis": analysis_result
                        }
                        st.session_state.history.insert(0, new_decision)
                        st.success("Аналіз успішно завершено!")
            else:
                st.warning("Будь ласка, введіть опис ситуації.")

    # Відображення поточного результату
    if st.session_state.history:
        current = st.session_state.history[0]
        
        st.divider()
        st.header(f"Результат для: {current['query']}")
        
        # Експорт
        docx_data = export_to_docx(current)
        st.download_button(
            label="📄 Скачати повний звіт (.docx)",
            data=docx_data,
            file_name=f"decision_analysis_{datetime.now().strftime('%d_%m')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # Tabs для зручності перегляду
        tab1, tab2, tab3 = st.tabs(["📊 Плюси/Мінуси", "📐 SWOT-аналіз", "📝 Підсумок"])
        
        with tab1:
            col1, col2 = st.columns(2)
            with col1:
                st.success("**Переваги**")
                for p in current['analysis']['prosCons']['pros']:
                    st.write(f"- {p}")
            with col2:
                st.error("**Недоліки**")
                for c in current['analysis']['prosCons']['cons']:
                    st.write(f"- {c}")
                    
        with tab2:
            s1, s2 = st.columns(2)
            s3, s4 = st.columns(2)
            with s1:
                st.info("**Strengths**\n" + "\n".join([f"- {x}" for x in current['analysis']['swot']['strengths']]))
            with s2:
                st.warning("**Weaknesses**\n" + "\n".join([f"- {x}" for x in current['analysis']['swot']['weaknesses']]))
            with s3:
                st.info("**Opportunities**\n" + "\n".join([f"- {x}" for x in current['analysis']['swot']['opportunities']]))
            with s4:
                st.error("**Threats**\n" + "\n".join([f"- {x}" for x in current['analysis']['swot']['threats']]))
                
        with tab3:
            st.markdown(f"### Рекомендація\n> {current['analysis']['summary']}")

# Бокова панель
st.sidebar.title("📚 Історія аналізів")
if not st.session_state.history:
    st.sidebar.write("Історія поки порожня.")
else:
    for i, item in enumerate(st.session_state.history):
        btn_label = f"{item['timestamp']}: {item['query'][:25]}..."
        if st.sidebar.button(btn_label, key=f"hist_{i}"):
            # Переміщуємо вибране на перше місце (трохи простовато, зате працює для Streamlit)
            st.session_state.history.insert(0, st.session_state.history.pop(i))
            st.rerun()
